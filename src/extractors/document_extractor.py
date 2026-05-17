"""DocumentExtractor — archivos locales (PDF, DOCX, EPUB, CSV, TXT, MD).

Fase 1: parsing cacheado por hash de archivo. Re-procesar el mismo PDF gigante
no vuelve a parsearlo desde cero.
"""
from __future__ import annotations

import asyncio
from pathlib import Path

from src.extractors.base import (
    Extractor,
    ExtractedContent,
    IngestionPayload,
    SourceKind,
)
from src.utils.cache import cache, hash_file
from src.utils.logger import get_logger

logger = get_logger(__name__)

SUPPORTED_EXTS = {
    ".pdf", ".docx", ".epub", ".csv", ".txt", ".md",
    # Phase 3 — OCR de imágenes
    ".jpg", ".jpeg", ".png", ".gif", ".bmp", ".tiff", ".tif", ".webp",
}


class DocumentExtractor(Extractor):
    name = "document"

    def can_handle(self, payload: IngestionPayload) -> bool:
        if payload.kind is not SourceKind.FILE:
            return False
        return Path(payload.raw).suffix.lower() in SUPPORTED_EXTS

    async def extract(self, payload: IngestionPayload) -> ExtractedContent:
        path = Path(payload.raw)
        ext = path.suffix.lower()
        original_name = payload.metadata.get("filename", path.name)
        logger.info("DocumentExtractor: %s (%s)", original_name, ext)

        # Cache: solo aplica a formatos pesados (PDF/EPUB/DOCX). Los .md/.txt/.csv son baratos.
        # Cachear formatos costosos (OCR y parsing pesado)
        _CACHED_EXTS = {".pdf", ".epub", ".docx", ".jpg", ".jpeg", ".png",
                        ".gif", ".bmp", ".tiff", ".tif", ".webp"}
        if ext in _CACHED_EXTS:
            file_hash = await asyncio.to_thread(hash_file, path)
            cached = cache.get(f"doc{ext}", file_hash)
            if cached is not None:
                text = cached
            else:
                text = await asyncio.to_thread(self._dispatch_parse, path, ext)
                cache.set(f"doc{ext}", file_hash, text)
        else:
            text = await asyncio.to_thread(self._dispatch_parse, path, ext)

        # Metadatos enriquecidos según tipo
        extra: dict = {"original_filename": original_name}
        if ext == ".pdf":
            acad = await asyncio.to_thread(self._extract_academic_metadata, path, text)
            extra.update(acad)
            extra["doc_kind"] = "paper" if "doi" in acad else "pdf"
        elif ext == ".epub":
            book_meta = await asyncio.to_thread(self._extract_epub_metadata, path)
            extra.update(book_meta)
            extra["doc_kind"] = "book"
        elif ext == ".docx":
            docx_meta = await asyncio.to_thread(self._extract_docx_metadata, path)
            extra.update(docx_meta)
            extra["doc_kind"] = "document"
        elif ext in {".jpg", ".jpeg", ".png", ".gif", ".bmp", ".tiff", ".tif", ".webp"}:
            extra["doc_kind"] = "image"
            from src.utils.vision import get_image_metadata
            extra.update(await asyncio.to_thread(get_image_metadata, path))
            # Si la imagen vino con caption desde Telegram, lo añadimos al texto
            # como contexto del usuario y como pista de título.
            caption = payload.metadata.get("caption")
            if caption:
                extra["caption"] = caption
                text = (
                    f"## Caption del usuario\n\n{caption}\n\n---\n\n{text}"
                )
        else:
            extra["doc_kind"] = "text"

        # Si la metadata trae un título mejor, lo usamos como title_hint
        title_hint = (
            extra.get("book_title")
            or extra.get("doc_title")
            or extra.get("caption")
            or Path(original_name).stem
        )
        # Caption muy largo → recortar para el title_hint
        if len(title_hint) > 80:
            title_hint = title_hint[:77].rstrip() + "…"

        return ExtractedContent(
            title_hint=title_hint,
            source_type=ext.lstrip("."),
            source_ref=original_name,
            text=text,
            extra=extra,
        )

    # --- Dispatcher síncrono ---
    def _dispatch_parse(self, path: Path, ext: str) -> str:
        match ext:
            case ".pdf":
                return self._parse_pdf(path)
            case ".docx":
                return self._parse_docx(path)
            case ".epub":
                return self._parse_epub(path)
            case ".csv":
                return self._parse_csv(path)
            case ".txt" | ".md":
                return path.read_text(encoding="utf-8", errors="ignore")
            case _ if ext in {".jpg", ".jpeg", ".png", ".gif",
                               ".bmp", ".tiff", ".tif", ".webp"}:
                return self._parse_image(path)
        raise ValueError(f"Extensión no soportada: {ext}")

    @staticmethod
    def _parse_pdf(path: Path) -> str:
        import pymupdf
        doc = pymupdf.open(path)
        try:
            pages = [page.get_text("text") for page in doc]
            text = "\n\n".join(pages)
            # Scanned PDF: fall back to OCR if almost no text was found
            if len(text.strip()) < 50:
                logger.info("PDF sin capa de texto, intentando OCR…")
                try:
                    ocr_pages = []
                    for page in doc:
                        tp = page.get_textpage_ocr(flags=0, full=True)
                        ocr_pages.append(page.get_text(textpage=tp))
                    text = "\n\n".join(ocr_pages)
                except Exception as ocr_err:
                    logger.warning("OCR falló: %s", ocr_err)
            return text
        finally:
            doc.close()

    @staticmethod
    def _parse_docx(path: Path) -> str:
        from docx import Document
        doc = Document(str(path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # también capturamos texto de tablas
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    paragraphs.append(" | ".join(cells))
        return "\n".join(paragraphs)

    @staticmethod
    def _parse_epub(path: Path) -> str:
        from bs4 import BeautifulSoup
        from ebooklib import epub, ITEM_DOCUMENT

        book = epub.read_epub(str(path))
        chunks: list[str] = []
        for item in book.get_items_of_type(ITEM_DOCUMENT):
            soup = BeautifulSoup(item.get_content(), "lxml")
            chunks.append(soup.get_text(separator="\n", strip=True))
        return "\n\n".join(chunks)

    @staticmethod
    def _parse_image(path: Path) -> str:
        """Comprensión de imagen: vision LLM (qué hay) + OCR Tesseract (texto exacto).

        El resultado es un Markdown estructurado con dos partes claramente
        separadas. Ese texto entra al pipeline de síntesis como cualquier otro
        contenido y produce la nota final con todas las secciones habituales.
        """
        # 1. OCR puro — siempre lo intentamos (es rápido y barato)
        ocr_text = DocumentExtractor._ocr_image(path)

        # 2. Descripción semántica con modelo de visión (puede no estar disponible)
        from src.utils.vision import describe_image
        vision_md = describe_image(path)

        if not vision_md and not ocr_text:
            return f"[Imagen sin texto legible ni descripción: {path.name}]"

        parts: list[str] = []
        if vision_md:
            parts.append("## Análisis visual\n\n" + vision_md.strip())
        if ocr_text:
            parts.append("## Texto extraído (OCR Tesseract)\n\n" + ocr_text.strip())
        return "\n\n---\n\n".join(parts)

    @staticmethod
    def _ocr_image(path: Path) -> str:
        """OCR puro vía PyMuPDF + Tesseract."""
        import pymupdf
        try:
            doc = pymupdf.open(str(path))
        except Exception:
            return ""
        try:
            parts: list[str] = []
            for page in doc:
                tp = page.get_textpage_ocr(flags=0, full=True)
                parts.append(page.get_text(textpage=tp))
            return "\n\n".join(parts).strip()
        except Exception:
            return ""
        finally:
            doc.close()

    # --- Extractores de metadata por formato ---

    @staticmethod
    def _extract_epub_metadata(path: Path) -> dict:
        """Lee Dublin Core de un EPUB (título, autor, editor, fecha, ISBN, descripción, tags)."""
        from ebooklib import epub
        meta: dict = {}
        try:
            book = epub.read_epub(str(path))
        except Exception:
            return meta

        def _first(field: str) -> str | None:
            try:
                items = book.get_metadata("DC", field)
                if items:
                    val = items[0][0]
                    return val.strip() if isinstance(val, str) else None
            except Exception:
                pass
            return None

        def _all(field: str) -> list[str]:
            try:
                items = book.get_metadata("DC", field)
                return [i[0].strip() for i in items if isinstance(i[0], str) and i[0].strip()]
            except Exception:
                return []

        if t := _first("title"):
            meta["book_title"] = t
        if authors := _all("creator"):
            meta["authors"] = authors[:10]
        if pub := _first("publisher"):
            meta["publisher"] = pub
        if date := _first("date"):
            meta["published"] = date[:10]  # YYYY o YYYY-MM-DD
            if len(date) >= 4 and date[:4].isdigit():
                meta["year"] = int(date[:4])
        if lang := _first("language"):
            meta["original_language"] = lang
        if desc := _first("description"):
            # Limpiar HTML simple
            from bs4 import BeautifulSoup
            clean = BeautifulSoup(desc, "lxml").get_text(separator=" ", strip=True)
            if clean:
                meta["description"] = clean[:600]
        # Identificadores: ISBN
        try:
            for ident in book.get_metadata("DC", "identifier"):
                val = ident[0] if isinstance(ident[0], str) else ""
                # Formato ISBN13 o ISBN10
                import re as _re
                m = _re.search(r"(97[89][\-\s]?\d{1,5}[\-\s]?\d{1,7}[\-\s]?\d{1,7}[\-\s]?\d|\d{9}[\dXx])", val)
                if m:
                    meta["isbn"] = m.group(1).replace(" ", "").replace("-", "")
                    break
        except Exception:
            pass
        # Tags / temas
        if subjects := _all("subject"):
            meta["source_tags"] = subjects[:10]
        return meta

    @staticmethod
    def _extract_docx_metadata(path: Path) -> dict:
        """Lee core_properties de un .docx (autor, fecha, título)."""
        from docx import Document
        meta: dict = {}
        try:
            doc = Document(str(path))
            cp = doc.core_properties
            if cp.title:
                meta["doc_title"] = cp.title.strip()
            if cp.author:
                meta["author"] = cp.author.strip()
            if cp.created:
                meta["published"] = cp.created.date().isoformat()
                meta["year"] = cp.created.year
            if cp.subject:
                meta["description"] = cp.subject.strip()[:600]
            if cp.keywords:
                tags = [t.strip() for t in cp.keywords.split(",") if t.strip()]
                if tags:
                    meta["source_tags"] = tags[:10]
        except Exception:
            pass
        return meta

    @staticmethod
    def _extract_academic_metadata(path: Path, text: str) -> dict:
        """Extrae DOI, autores y abstract de PDFs académicos (Phase 3)."""
        import re
        import pymupdf
        meta: dict = {}

        # DOI
        doi_match = re.search(r"10\.\d{4,9}/[^\s\]\)>\",']+", text)
        if doi_match:
            meta["doi"] = doi_match.group(0).rstrip(".,")

        # Abstract (primeros 3 000 chars del texto)
        abs_match = re.search(
            r"Abstract[:\s]*\n(.+?)(?=\n\n|\n#{1,3}|\nIntroduction|\n1[\.\s])",
            text[:3000],
            re.DOTALL | re.IGNORECASE,
        )
        if abs_match:
            meta["abstract"] = abs_match.group(1).strip()[:500]

        # Autores desde metadatos del PDF
        try:
            doc = pymupdf.open(str(path))
            pdf_meta = doc.metadata or {}
            doc.close()
            raw_author = pdf_meta.get("author", "").strip()
            if raw_author:
                authors = [a.strip() for a in re.split(r"[;,]", raw_author) if a.strip()]
                if authors:
                    meta["authors"] = authors[:10]
        except Exception:
            pass

        return meta

    @staticmethod
    def _parse_csv(path: Path) -> str:
        import pandas as pd
        df = pd.read_csv(path)
        # Markdown table — el LLM lo entiende bien
        head = df.head(50).to_markdown(index=False)
        meta = (
            f"_CSV con {len(df)} filas y {len(df.columns)} columnas. "
            f"Columnas: {', '.join(df.columns)}_\n\n"
        )
        return meta + head
